#This program was revised on May, 19th, 2019, being planned to launch on May, 21st, 2019.
#Version: 1.3
#Copyright: Luo Ziqian
'''README
	This program is for people who need to extract a certain part of a file in a regular way and only available to .txt/.epud
	For instance: Using this program to process novels
'''
import re
import os
import winreg
import tkinter
import tkinter.filedialog
import tkinter.messagebox
import tkinter.font
import cn2an
import docx


#file operations
def getForm(file):
    pattern = r'(\.txt)|(\.docx)|(\.epud)|(\.mobi)'
    return re.search(pattern, file).group()


def getOutputFile(file, start, end):
    parts = file.split('/') #split apart 
    parts[-1] = start + '-' + end + '_' + parts[-1] #change the last part of a file path
    return '/'.join(parts) #rejoin them


def getFileSize(filePath):
    fsize = os.path.getsize(filePath)
    fsize = fsize / float(1024 * 1024)
    return round(fsize, 2)


def getDesktop():
    key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders')
    return winreg.QueryValueEx(key, "Desktop")[0]


def selectFile(file):
    f = tkinter.filedialog.askopenfilename(title=u'选择文件', initialdir=(os.path.expanduser(getDesktop())))
    return file.set(f)


#read&write
def read(file, form): #assign different forms of files to different functions. Now .epud and .mobi are not available
    if form == '.txt':
        return readInTxt(file)
    elif form == '.docx':
        return readInDocx(file)
    elif form == '.epud':
        return readInEpud(file)
    elif form == '.mobi':
        return readInMobi(file)
    else:
        return None


def readInTxt(file): 
    decodeList = ['utf-8', 'gb18030', 'gbk', 'gb2312', 'ISO-8859-2', 'Error'] #predicted encoding list 
    for decode in decodeList: #try them one by one
        try:
            with open(file, 'r', encoding=decode) as f:
                return f.read()
        except:
            if decode == 'Error':
                raise Exception('Decoding Failed')
            continue


def readInDocx(file):
    doc = docx.Document(file)
    text = []
    paras = doc.paragraphs
    for element in paras:
        text.append(element.text)
    return '\n'.join(text)
    

def readInEpud(file):
    pass


def readInMobi(file):
    pass


def write(file, content, form): #assign different forms of files to different functions
    if form == '.txt':
        return writeInTxt(file, content)
    elif form == '.docx':
        return writeInDocx(file, content)
    elif form == '.epud':
        return writeInEpud(file, content)
    elif form == '.mobi':
        return writeInMobi(file, content)
    else:
        return None
        

def writeInTxt(file, content):
    decodeList = ['gbk', 'gb18030', 'gb2312', 'utf-8', 'ISO-8859-2', 'Error'] #predicted encoding list 
    for decode in decodeList: #try them one by one
        try:
            with open(file, 'wb') as f:
                f.write(content.encode(decode))
                break
        except:
            if decode == 'Error':
                raise Exception('Decoding Failed')
            continue


def writeInDocx(file, content):
    doc = docx.Document()
    doc.add_paragraph(content)
    doc.save(file)


def writeInEpud(file, content):
    pass


def writeInMobi(file, content):
    pass


#Now this function is not available
def openFile(file, form):
    if form == '.txt':
        os.system('notepad ' + file)
    elif form == '.docx':
        os.system('' + file)
    elif form == '.epud':
        os.system('' + file)
    elif form == 'mobi':
        os.system('' + file)


#copy file
def combine(string):
    return '第' + string + '章'


def duplicate(string): #convert to different expressions
    return combine(string), combine(cn2an.an2cn(int(string)))


def pattern(string, start, end): #assign a proper pattern
    if end == '-1':  
        start, _start = duplicate(start)
        return '(({start})|({_start}))(.+)'.format(start = start, _start = _start)
    else:
        start, _start = duplicate(start)
        end, _end = duplicate(str(int(end) + 1))
        return '(({start})|({_start}))(.+?)(({end})|({_end}))'.format(start = start, _start = _start, end = end, _end = _end)


def textCopy(file, start, end, window):
    try:
        form = getForm(file) #get the form of the document
        content = read(file, form) #get content
        try:
            reExp = pattern(content, start, end) #get regular expression
            result = re.search(reExp, content, re.DOTALL).group() #get the processed content
            outputFile = getOutputFile(file, start, end) #get the output file name
            write(outputFile, result, form) #write processed content into the output file
            temp = tkinter.messagebox.askyesno(title="Success", message='File size: ' + str(getFileSize(outputFile)) + 'MB (:\n' + 'Open file?') #try to ask a further question
            if temp:
                os.system('notepad '+outputFile) #open the file in order to check whether the content is wanted
            else:
                window.quit() #otherwise quit the program
        except AttributeError: #if the content is None, a messagebox will be poped
            temp = tkinter.messagebox.askretrycancel(title="Fail", message='The Clip does not exist! ):') 
            if temp == False:
                window.quit()
    except Exception as e: #show the errors
        tkinter.messagebox.showerror(message='Error: ' + str(e))


#GUI
def generateWindow(title="", geometry="600x150"): #produce a window
    window = tkinter.Tk()
    window.title(title)
    window.geometry(geometry)
    return window


def main():
    mainWindow = generateWindow(title="TXT Scissor", geometry="300x90")
    #some variables
    file = tkinter.StringVar(mainWindow)
    start = tkinter.StringVar(mainWindow, value='1')
    end = tkinter.StringVar(mainWindow, value='-1')
    font = tkinter.font.Font(family='Arial', size=12, weight=tkinter.font.NORMAL)

    #define components
    label1 = tkinter.Label(mainWindow, text='Initial Value: ', width=15, font=font)
    label2 = tkinter.Label(mainWindow, text='End Value: ', width=15, font=font)

    entry1 = tkinter.Entry(mainWindow, textvariable=start, width=15, font=font)
    entry2 = tkinter.Entry(mainWindow, textvariable=end, width=15, font=font)

    selectButton = tkinter.Button(mainWindow, text=' Select ', command=lambda: selectFile(file), width=10, font=font)
    startButton = tkinter.Button(mainWindow, text=' Start ', command=lambda: textCopy(file.get(), start.get(),
                                                    end.get(), mainWindow), width=10, font=font)

    #grid the components
    label1.grid(row=0, column=0)
    entry1.grid(row=0, column=1)
    label2.grid(row=1, column=0)
    entry2.grid(row=1, column=1)
    selectButton.grid(row=2, column=0)
    startButton.grid(row=2, column=1)

    mainWindow.mainloop()
    
	
#test(only for Bi's Device)
def test():
    window = generateWindow()
    textCopy('D:/test files for txtScissor and others/三界淘宝店.txt', '99', '-1', window)
    textCopy('D:/test files for txtScissor and others/三界淘宝店.txt', '3', '334', window)
    textCopy('D:/test files for txtScissor and others/三界淘宝店.docx', '6', '334', window)
    textCopy('D:/test files for txtScissor and others/三界淘宝店.docx', '443', '-1', window)
    textCopy('D:/test files for txtScissor and others/修真高手在校园.txt', '341', '442', window)


if __name__ == '__main__':
    # test()
    main()
    








